"""
Generate a senior-lawyer-quality benchmark legal notice and save to output/.

This serves as a reference to measure the drafting agent's output quality.
Scenario: Recovery of ₹8,50,000 hand loan, Bengaluru, bank transfer proof only, 18% interest.
"""
import os
import sys
import re
from datetime import datetime

# Fix Windows console encoding
if sys.stdout.encoding != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8")

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── Senior Lawyer Benchmark Draft ──────────────────────────────────────
# This is what a well-drafted legal notice should look like.
# The agent should produce output of THIS quality.

BENCHMARK_DRAFT = """LEGAL NOTICE
(Under Section 80 of the Code of Civil Procedure, 1908)

Date: {{NOTICE_DATE}}

THROUGH REGISTERED POST A.D. / SPEED POST

TO,

{{BORROWER_NAME}},
{{BORROWER_ADDRESS}},
Bengaluru, Karnataka – {{BORROWER_PIN}}.

FROM,

{{LENDER_NAME}},
{{LENDER_ADDRESS}},
Bengaluru, Karnataka – {{LENDER_PIN}}.

Through: {{ADVOCATE_NAME}}, Advocate
{{ADVOCATE_ADDRESS}}, Bengaluru
Enrolment No: {{ADVOCATE_ENROLLMENT_NO}}

SUBJECT: LEGAL NOTICE DEMANDING REPAYMENT OF HAND LOAN OF ₹8,50,000/- (RUPEES EIGHT LAKHS FIFTY THOUSAND ONLY) ALONG WITH INTEREST AT 18% PER ANNUM

Sir/Madam,

Under the instructions from and on behalf of my client, {{LENDER_NAME}} (hereinafter referred to as "the Noticee"), I do hereby serve upon you this Legal Notice and state as follows:

1. That my client is a law-abiding citizen residing at {{LENDER_ADDRESS}}, Bengaluru, Karnataka, and is well known to you personally.

2. That you approached my client on or about {{LOAN_DATE}} and requested for a hand loan of ₹8,50,000/- (Rupees Eight Lakhs Fifty Thousand Only) citing urgent personal/financial necessity. You made solemn assurances and promises that you would repay the entire amount within a period of {{REPAYMENT_PERIOD}} from the date of advancement of the loan.

3. That based on the trust and confidence reposed in you and relying upon your verbal assurances, my client advanced a sum of ₹8,50,000/- (Rupees Eight Lakhs Fifty Thousand Only) to you by way of electronic bank transfer on {{TRANSFER_DATE}} from my client's bank account maintained at {{LENDER_BANK_NAME}}, Account No. {{LENDER_ACCOUNT_NUMBER}} to your bank account maintained at {{BORROWER_BANK_NAME}}, Account No. {{BORROWER_ACCOUNT_NUMBER}}, bearing UTR/Transaction Reference No. {{TRANSACTION_REFERENCE}}.

4. That the said bank transfer statement/receipt is in the possession of my client and shall serve as conclusive proof of the said transaction and advancement of the loan amount.

5. That although no written loan agreement was executed between the parties, the bank transfer records, electronic communications, and the relationship of trust between the parties clearly establish the existence of the hand loan and the obligation on your part to repay the same.

6. That despite the expiry of the agreed repayment period and despite several verbal and telephonic reminders and demands made by my client, you have wilfully, deliberately, and without any justifiable reason failed, neglected, and refused to repay the loan amount of ₹8,50,000/- or any part thereof to my client.

7. That your failure to repay the said amount despite repeated demands amounts to:
   (a) Civil breach of the obligation to repay a debt;
   (b) Criminal breach of trust punishable under Section 405 read with Section 406 of the Bharatiya Nyaya Sanhita, 2023 (erstwhile Section 405/406 of the Indian Penal Code, 1860);
   (c) Cheating punishable under Section 318 of the Bharatiya Nyaya Sanhita, 2023 (erstwhile Section 420 of the Indian Penal Code, 1860);
   (d) Dishonest misappropriation of property under the applicable provisions of law.

8. That my client is entitled to claim interest at the rate of 18% per annum on the principal amount of ₹8,50,000/- from the date of advancement of the loan, i.e., {{LOAN_DATE}}, till the date of actual realization. The accrued interest as on the date of this notice is approximately ₹{{ACCRUED_INTEREST}}/-.

9. That my client has suffered immense mental agony, harassment, and financial hardship due to your deliberate failure to honour your commitment to repay the said loan.

DEMAND:

10. In view of the above facts and circumstances, you are hereby called upon and demanded to pay to my client the following amounts within a period of 15 (FIFTEEN) days from the date of receipt of this notice:

   (i)   Principal loan amount:                          ₹8,50,000/-
   (ii)  Interest at 18% per annum till date:            ₹{{ACCRUED_INTEREST}}/-
   (iii) Cost of this Legal Notice:                      ₹{{NOTICE_COST}}/-
         ──────────────────────────────────────────────
   TOTAL:                                                ₹{{TOTAL_AMOUNT}}/-

   (Rupees {{TOTAL_AMOUNT_WORDS}} Only)

11. The above-mentioned amount shall be paid by way of RTGS/NEFT bank transfer to the following account:

   Account Name:   {{LENDER_NAME}}
   Bank:           {{LENDER_BANK_NAME}}
   Account No:     {{LENDER_ACCOUNT_NUMBER}}
   IFSC Code:      {{LENDER_IFSC_CODE}}

CONSEQUENCES:

12. Please take note that in the event of your failure to comply with this notice and make the payment of the demanded amount within the stipulated period of 15 days, my client shall be left with no alternative but to initiate:

   (a) Civil recovery proceedings before the competent Civil Court/Consumer Forum at Bengaluru for recovery of the entire amount along with pendente lite and future interest;
   (b) Criminal proceedings under the relevant provisions of the Bharatiya Nyaya Sanhita, 2023 and other applicable laws;
   (c) Such other legal proceedings as may be advised,

   all at your sole risk, cost, and consequences, and my client shall also claim the costs of litigation, advocate fees, and all incidental expenses from you.

13. This notice is issued without prejudice to any and all other rights and remedies available to my client under law, all of which are expressly reserved.

14. A copy of this notice is being retained in my office for record and reference.

VERIFICATION

I, {{LENDER_NAME}}, the sender/client mentioned herein, do hereby state and verify that the contents of this Legal Notice are true and correct to the best of my knowledge, information, and belief. No part hereof is false and nothing material has been concealed.

Verified at Bengaluru, Karnataka, on this {{NOTICE_DATE}}.


_________________________________
{{LENDER_NAME}}
(Sender / Client)


Drafted, Settled & Issued by:

_________________________________
{{ADVOCATE_NAME}}
Advocate, High Court of Karnataka
Enrolment No: {{ADVOCATE_ENROLLMENT_NO}}
{{ADVOCATE_ADDRESS}}
Bengaluru, Karnataka
Mobile: {{ADVOCATE_MOBILE}}
Email: {{ADVOCATE_EMAIL}}

LIST OF ENCLOSURES:

1. Certified copy of bank transfer statement/receipt dated {{TRANSFER_DATE}} evidencing the loan of ₹8,50,000/-
2. Copy of identification document of the sender
3. Record of telephonic/electronic communication demanding repayment (if available)
4. Any other relevant documents supporting the claim

MODE OF SERVICE:

This Legal Notice is being sent via:
1. Registered Post with Acknowledgement Due (RPAD) to the above-mentioned address
2. Speed Post to the above-mentioned address
3. Email at {{BORROWER_EMAIL}} (if available)

Place: Bengaluru
Date: {{NOTICE_DATE}}"""


def save_draft(text, name):
    """Save as TXT and DOCX."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base = f"{name}_{timestamp}"
    saved = {}

    # TXT
    txt_path = os.path.join(OUTPUT_DIR, f"{base}.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(text)
    saved["txt"] = txt_path

    # DOCX
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1)

        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
            elif stripped.isupper() and len(stripped) < 100:
                h = doc.add_heading(stripped, level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif stripped.startswith("SUBJECT:"):
                p = doc.add_paragraph()
                r = p.add_run(stripped)
                r.bold = True
                r.font.size = Pt(11)
            elif re.match(r"^\d+\.", stripped):
                doc.add_paragraph(stripped)
            elif stripped.startswith("(") and stripped.endswith(")"):
                p = doc.add_paragraph(stripped)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif "___" in stripped:
                doc.add_paragraph(stripped)
            else:
                doc.add_paragraph(stripped)

        docx_path = os.path.join(OUTPUT_DIR, f"{base}.docx")
        doc.save(docx_path)
        saved["docx"] = docx_path
    except ImportError:
        pass

    return saved


def analyze_quality(text):
    """Analyze the draft quality metrics."""
    lines = text.strip().split("\n")
    words = text.split()
    placeholders = re.findall(r"\{\{([A-Za-z_]+)\}\}", text)
    unique_ph = sorted(set(placeholders))
    paragraphs = [l for l in lines if l.strip() and re.match(r"^\d+\.", l.strip())]
    sections = []
    for keyword in ["DEMAND:", "CONSEQUENCES:", "VERIFICATION", "LIST OF ENCLOSURES:", "MODE OF SERVICE:"]:
        if keyword in text:
            sections.append(keyword.rstrip(":"))

    print("\n── QUALITY ANALYSIS ──")
    print(f"  Word count:          {len(words)}")
    print(f"  Line count:          {len(lines)}")
    print(f"  Numbered paragraphs: {len(paragraphs)}")
    print(f"  Sections found:      {', '.join(sections)}")
    print(f"  Placeholders:        {len(placeholders)} total, {len(unique_ph)} unique")
    print(f"\n  Key legal elements:")

    checks = {
        "Section 80 CPC reference": "Section 80" in text,
        "BNS/IPC criminal sections": "405" in text or "318" in text,
        "Interest rate stated": "18%" in text,
        "Principal amount stated": "8,50,000" in text,
        "15-day notice period": "15" in text and "days" in text.lower(),
        "Verification clause": "VERIFICATION" in text,
        "Advocate endorsement": "ADVOCATE" in text.upper(),
        "Enclosures listed": "ENCLOSURE" in text.upper(),
        "Mode of service": "MODE OF SERVICE" in text.upper(),
        "RPAD mentioned": "RPAD" in text or "Registered Post" in text,
        "Bank transfer proof": "bank transfer" in text.lower(),
        "Demand section": "DEMAND" in text,
        "Consequences section": "CONSEQUENCES" in text or "COMPLIANCE" in text.upper(),
    }

    score = 0
    for check, passed in checks.items():
        status = "PASS" if passed else "FAIL"
        print(f"    [{status}] {check}")
        if passed:
            score += 1

    print(f"\n  Quality score: {score}/{len(checks)} ({round(score/len(checks)*100)}%)")
    return score, len(checks)


def main():
    print("=" * 60)
    print("BENCHMARK: Senior Lawyer Legal Notice")
    print("Scenario: ₹8,50,000 hand loan, Bengaluru, bank transfer proof")
    print("=" * 60)

    saved = save_draft(BENCHMARK_DRAFT, "benchmark_legal_notice")

    print("\nSaved files:")
    for fmt, path in saved.items():
        size = os.path.getsize(path)
        print(f"  {fmt.upper()}: {os.path.basename(path)} ({size:,} bytes)")

    analyze_quality(BENCHMARK_DRAFT)

    print("\n── WHAT MAKES THIS DRAFT GOOD ──")
    print("  1. Proper header: 'LEGAL NOTICE (Under Section 80 CPC)'")
    print("  2. Correct addressing format (TO/FROM with advocate)")
    print("  3. Clear subject line with amount in words")
    print("  4. Numbered paragraphs with logical flow")
    print("  5. Criminal sections cited (BNS 2023, not old IPC)")
    print("  6. Interest calculation methodology stated")
    print("  7. Clear DEMAND section with itemized amounts")
    print("  8. Payment details (bank account for repayment)")
    print("  9. CONSEQUENCES section with civil + criminal warnings")
    print("  10. Proper verification clause (Karnataka format)")
    print("  11. Advocate signature block with enrollment number")
    print("  12. Enclosures list")
    print("  13. Mode of service (RPAD + Speed Post + Email)")
    print("  14. Placeholders for ALL unknown info — nothing fabricated")

    print("\n── PROMPT IMPROVEMENT AREAS ──")
    print("  The drafting agent needs prompt guidance on:")
    print("  1. STRUCTURE: Follow a logical section order (facts → demand → consequences)")
    print("  2. COMPLETENESS: Include ALL required sections (verification, enclosures, service)")
    print("  3. CITATIONS: Reference correct current law (BNS 2023, not old IPC)")
    print("  4. FORMALITY: Use proper legal language and numbering")
    print("  5. DEMAND FORMAT: Itemize amounts in a clear table")
    print("  6. ADVOCATE BLOCK: Include proper signature and enrollment")
    print("  Done!")


if __name__ == "__main__":
    main()
