"""
Legal Drafting Assistant - Chat-style UI (stateless, no local history storage).

Usage:
  1) Start LangGraph server:  langgraph dev
  2) Run app:                  python -m streamlit run research/drafting_app.py
"""

from __future__ import annotations

import io
import json
import re
import time
import requests
import streamlit as st

LANGGRAPH_URL = "http://localhost:2024"
GRAPH_NAME = "legal_drafting_agent"

st.set_page_config(page_title="Legal Drafting Assistant", layout="wide")


def _inject_css() -> None:
    st.markdown(
        """
        <style>
        .block-container {
            max-width: 920px;
            padding-top: 1.5rem;
            padding-bottom: 2rem;
        }
        .title {
            font-size: 1.75rem;
            font-weight: 700;
            margin-bottom: 0.2rem;
        }
        .subtitle {
            color: #6b7280;
            margin-bottom: 1rem;
        }
        .draft-box {
            border: 1px solid #e5e7eb;
            border-radius: 14px;
            padding: 1rem;
            background: #ffffff;
        }
        .note {
            color: #6b7280;
            font-size: 0.9rem;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def _url() -> str:
    return st.session_state.get("langgraph_url", LANGGRAPH_URL)


def _create_thread() -> str:
    # Stateless behavior: create a fresh thread for each request.
    resp = requests.post(
        f"{_url()}/threads",
        json={},
        headers={"Content-Type": "application/json"},
        timeout=10,
    )
    resp.raise_for_status()
    return resp.json()["thread_id"]


def _normalize_text(text: str) -> str:
    if not isinstance(text, str):
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)
    return text.strip()


def _is_internal_control_message(text: str) -> bool:
    if not isinstance(text, str):
        return False
    t = text.strip()
    if not t:
        return False

    # Known pipeline control prefixes.
    if t.startswith("REVIEW_INPUT_JSON") or t.startswith("WORKFLOW_STATE_JSON"):
        return True

    # Raw JSON control payloads accidentally surfaced as assistant text.
    if t[:1] in {"{", "["}:
        try:
            parsed = json.loads(t)
            if isinstance(parsed, dict):
                control_keys = {"review_pass", "draft_context", "workflow_state", "parallel_outputs"}
                if any(k in parsed for k in control_keys):
                    return True
        except Exception:
            pass

    return False


def _extract_backend_error(payload) -> str:
    if not isinstance(payload, dict):
        return ""

    err = payload.get("__error__")
    if err:
        if isinstance(err, dict):
            err_type = _normalize_text(str(err.get("error") or err.get("type") or ""))
            message = _normalize_text(str(err.get("message") or ""))
            if err_type and message and err_type.lower() not in message.lower():
                return f"{err_type}: {message}"
            return message or err_type
        return _normalize_text(str(err))

    for key in ("error_message", "error"):
        value = payload.get(key)
        if isinstance(value, str) and value.strip():
            return _normalize_text(value)

    nested = payload.get("values")
    if isinstance(nested, dict):
        return _extract_backend_error(nested)

    return ""


def _extract_blob_text(blob) -> str:
    if isinstance(blob, bytes):
        return blob.decode("utf-8", errors="ignore")
    if isinstance(blob, str):
        text = _normalize_text(blob)
        # Tool messages often return JSON strings; parse and extract meaningful text.
        if text and text[:1] in {"{", "["}:
            try:
                parsed = json.loads(text)
                nested = _extract_blob_text(parsed)
                if nested:
                    return nested
            except Exception:
                pass
        return text
    if isinstance(blob, list):
        candidates = [_extract_blob_text(x) for x in blob]
        candidates = [c for c in candidates if isinstance(c, str) and c.strip()]
        if not candidates:
            return ""
        return max(candidates, key=len)
    if isinstance(blob, dict):
        preferred_keys = (
            "export_content",
            "draft_content",
            "final_text",
            "final_draft",
            "draft_text",
            "text_content",
            "content",
            "text",
            "body",
            "document",
            "output",
            "data",
            "message",
        )
        for key in preferred_keys:
            txt = _extract_blob_text(blob.get(key))
            if txt:
                return txt
        # Generic fallback for unknown dict shapes.
        candidates = [_extract_blob_text(v) for v in blob.values()]
        candidates = [c for c in candidates if isinstance(c, str) and c.strip()]
        if candidates:
            return max(candidates, key=len)
    return ""


def _strip_draft_meta_preamble(text: str) -> str:
    if not isinstance(text, str):
        return ""

    if _is_internal_control_message(text):
        return ""

    lines = text.splitlines()
    if not lines:
        return ""

    meta_patterns = [
        r"(?i)^#+\s*.*\b(generating|drafting|final quality review|legal review)\b",
        r"(?i)^\s*(now|let me|i will|i'll)\b.*\b(draft|generate|review|retrieve|fetch|perform)\b",
        r"(?i)^\s*(template pack generated successfully|classification complete)\b",
    ]

    cleaned = []
    for idx, line in enumerate(lines):
        stripped = line.strip()
        if idx < 20:
            if any(re.search(p, stripped) for p in meta_patterns):
                continue
            if stripped in {"---", "___"} and idx < 8:
                continue
        cleaned.append(line)

    # Remove accidental triple blank runs after line filtering.
    out = []
    blank_run = 0
    for line in cleaned:
        if not line.strip():
            blank_run += 1
            if blank_run > 2:
                continue
        else:
            blank_run = 0
        out.append(line)

    return _normalize_text("\n".join(out))


def _looks_like_draft(text: str) -> bool:
    if not isinstance(text, str):
        return False
    if _is_internal_control_message(text):
        return False
    t = _strip_draft_meta_preamble(text)
    # Be format-agnostic: accept any substantial non-meta text.
    if len(t) < 80:
        return False
    low = t.lower()
    # Reject only clear pipeline/status chatter.
    meta_markers = [
        "document ready for use",
        "quick reference",
        "need help with something else",
        "your legal notice has been fully drafted",
        "provided above",
        "would you like me to proceed",
        "let me know and i'll proceed",
        "i already have your drafting_session_id",
        "i can see from the classification",
        "template pack generated successfully",
        "classification complete",
        "status:",
        "next step:",
        "ready for document drafting phase",
        "template summary",
        "critical missing info",
    ]
    # Allow long responses that include a short meta line plus the full draft.
    if any(m in low for m in meta_markers) and len(t) < 600:
        return False

    if _is_meta_narration(t):
        return False
    if _looks_like_review_report(t):
        return False

    # Minimal content sanity check.
    alpha_chars = sum(1 for ch in t if ch.isalpha())
    return alpha_chars >= 40


def _is_meta_narration(text: str) -> bool:
    if not isinstance(text, str):
        return False
    t = text.strip().lower()
    if not t:
        return False
    patterns = [
        r"^now\s+i(?:'| )?ll\b.*\b(draft|review|retrieve|fetch|perform|update|revise|correct)\b",
        r"^now\s+let\s+me\b.*\b(draft|review|retrieve|fetch|perform|update|revise|correct)\b",
        r"^let\s+me\b.*\b(draft|review|retrieve|fetch|perform|update|revise|correct)\b",
        r"^i\s+will\b.*\b(draft|review|retrieve|fetch|perform|update|revise|correct)\b",
        r"^i\s+need\s+to\b.*\b(update|revise|correct|fix)\b",
        r"\bvalidation\s+issues?\b",
        r"\brequired\s+corrections?\b",
        r"\bquality\s+review\b",
    ]
    return any(re.search(p, t) for p in patterns)


def _looks_like_review_report(text: str) -> bool:
    if not isinstance(text, str):
        return False
    low = text.strip().lower()
    if not low:
        return False
    markers = [
        "document review outcome",
        "structural completeness and readiness",
        "recommendation and next steps",
        "quality concerns and limitations",
        "action items for the attorney",
        "would you like me to proceed",
    ]
    hits = sum(1 for marker in markers if marker in low)
    return hits >= 2


def _is_agent_json_blob(text: str) -> bool:
    """Return True if the text is a raw JSON payload from an agent tool output.

    Agent nodes (classification, compliance, citation, localization, template, etc.)
    sometimes surface their structured JSON output directly as an assistant message.
    Legal drafts are never valid JSON, so any message whose entire content parses as
    a JSON dict or list is definitively an agent output blob and must not be shown in
    the chat bubble or treated as the draft.
    """
    if not isinstance(text, str):
        return False
    t = text.strip()
    if not t or t[0] not in {"{", "["}:
        return False
    try:
        parsed = json.loads(t)
        return isinstance(parsed, (dict, list))
    except Exception:
        return False


def _is_good_assistant_reply(text: str) -> bool:
    """Assistant chat bubble content should avoid internal review/report dumps."""
    if not isinstance(text, str):
        return False
    t = _normalize_text(text)
    if not t:
        return False
    if _is_internal_control_message(t):
        return False
    if _is_agent_json_blob(t):
        return False
    if _is_meta_narration(t):
        return False
    if _looks_like_review_report(t):
        return False
    return True


def _make_docx_bytes(draft_text: str) -> bytes | None:
    try:
        from docx import Document
        from docx.shared import Inches
    except Exception:
        return None

    draft_text = _strip_draft_meta_preamble(draft_text)
    if not draft_text:
        return None

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)

    for line in draft_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("###"):
            h = stripped.lstrip("#").strip()
            if h:
                doc.add_heading(h, level=3)
            continue
        if stripped.startswith("##"):
            h = stripped.lstrip("#").strip()
            if h:
                doc.add_heading(h, level=2)
            continue
        if stripped.startswith("#"):
            h = stripped.lstrip("#").strip()
            if h:
                doc.add_heading(h, level=1)
            continue
        if stripped.isupper() and 4 < len(stripped) < 100:
            doc.add_heading(stripped, level=2)
            continue
        if re.match(r"^={3,}|^-{3,}", stripped):
            continue
        doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _format_for_markdown_display(draft_text: str) -> str:
    """
    Improve heading rendering for legal drafts.
    - Converts all-caps short lines to markdown headings
    - Converts numbered section titles to subheadings when appropriate
    """
    text = _strip_draft_meta_preamble(draft_text)
    if not text:
        return ""

    out: list[str] = []
    lines = text.split("\n")

    for i, line in enumerate(lines):
        s = line.strip()
        if not s:
            out.append("")
            continue

        # Keep authored markdown headings as-is.
        if s.startswith("#"):
            out.append(s)
            continue

        # Convert short all-caps lines into headings.
        if s.isupper() and len(s) <= 120:
            level = "##" if i == 0 else "###"
            out.append(f"{level} {s}")
            continue

        # Promote likely section-title lines like "1. FACTS" / "2. JURISDICTION".
        if re.match(r"^\d+\.\s+[A-Z][A-Z\s/&,-]{3,}$", s):
            out.append(f"### {s}")
            continue

        out.append(s)

    return "\n".join(out)


def _parse_state(state: dict) -> dict:
    if isinstance(state, dict) and isinstance(state.get("values"), dict):
        state = state["values"]

    result = {
        "draft": "",
        "assistant_message": "",
        "recommendations": [],
        "doc_type": state.get("doc_type") or state.get("document_type", "draft"),
        "error": _extract_backend_error(state) or state.get("error_message"),
    }

    # Strong path: if export gate passed, trust and display its content as-is.
    export_output = state.get("export_output") or {}
    if isinstance(export_output, dict) and export_output.get("passed"):
        export_raw = _normalize_text(_extract_blob_text(export_output))
        if export_raw:
            result["draft"] = export_raw

    for key in ("draft_text", "export_output", "final_draft", "draft_v1"):
        if result["draft"]:
            break
        text = _strip_draft_meta_preamble(_extract_blob_text(state.get(key) or {}))
        if _looks_like_draft(text):
            result["draft"] = text
            break

    if not result["draft"]:
        fd = state.get("final_draft") or {}
        if isinstance(fd, dict):
            sections = fd.get("sections") or []
            if isinstance(sections, list):
                chunks = []
                for sec in sections:
                    if isinstance(sec, dict):
                        title = (sec.get("title") or "").strip()
                        content = _normalize_text(_extract_blob_text(sec.get("content", "")))
                        if title:
                            chunks.append(title)
                        if content:
                            chunks.append(content)
                candidate = _strip_draft_meta_preamble("\n\n".join(chunks))
                if _looks_like_draft(candidate):
                    result["draft"] = candidate

    for q in state.get("clarification_questions") or []:
        if isinstance(q, dict):
            field = q.get("field", "")
            question = q.get("question", q.get("text", ""))
            if field or question:
                result["recommendations"].append(f"{field}: {question}" if field else question)

    ai_messages = []
    tool_draft_candidates = []
    for msg in state.get("messages", []):
        if isinstance(msg, dict):
            role = str(msg.get("role") or msg.get("type", "")).lower()
            content_raw = _normalize_text(_extract_blob_text(msg.get("content", "")))
            if _is_internal_control_message(content_raw):
                continue

            name = str(msg.get("name") or msg.get("tool_name") or "").lower()
            content = _strip_draft_meta_preamble(content_raw)
            if role in ("assistant", "ai") and content and len(content) > 10 and _is_good_assistant_reply(content):
                ai_messages.append(content)
            if role in ("tool", "tool_message") or name in {
                "finalize_and_deliver",
                "get_draft_content",
                "save_draft",
            }:
                if _looks_like_draft(content):
                    tool_draft_candidates.append(content)

    if ai_messages:
        result["assistant_message"] = ai_messages[-1]
        if not result["draft"]:
            draft_candidates = [msg for msg in ai_messages if _looks_like_draft(msg)]
            if draft_candidates:
                # Prefer the latest draft-looking assistant output.
                best = draft_candidates[-1]
                result["draft"] = best
                if best == ai_messages[-1] and len(ai_messages) > 1:
                    result["assistant_message"] = ai_messages[-2]
    elif not result["assistant_message"] and tool_draft_candidates:
        result["assistant_message"] = "Final draft generated successfully."

    if not result["draft"] and tool_draft_candidates:
        result["draft"] = max(tool_draft_candidates, key=len)

    # Avoid rendering full draft twice (assistant bubble + Draft panel).
    # Use _looks_like_draft() instead of exact equality — minor whitespace differences
    # between draft_v1 content and the last AI message would otherwise cause both to render.
    if result["draft"] and _looks_like_draft(result.get("assistant_message", "")):
        result["assistant_message"] = "Final draft generated. See the Draft section below."

    # If we have a draft but reply text is empty/noisy, show a clean status line.
    if result["draft"] and not _is_good_assistant_reply(result.get("assistant_message", "")):
        result["assistant_message"] = "Final draft generated. See the Draft section below."

    return result


def _with_elapsed(payload: dict, started_at: float) -> dict:
    out = dict(payload or {})
    out["elapsed_seconds"] = round(time.perf_counter() - started_at, 3)
    return out


def send_message(user_message: str) -> dict:
    started_at = time.perf_counter()
    try:
        thread_id = _create_thread()
    except Exception as e:
        return _with_elapsed({"error": f"Cannot create thread: {e}"}, started_at)

    payload = {
        "assistant_id": GRAPH_NAME,
        "input": {"messages": [{"type": "human", "content": user_message}]},
    }
    try:
        resp = requests.post(
            f"{_url()}/threads/{thread_id}/runs/wait",
            json=payload,
            headers={"Content-Type": "application/json"},
            timeout=600,
        )
        if resp.status_code != 200:
            return _with_elapsed({"error": f"Server error {resp.status_code}: {resp.text[:300]}"}, started_at)

        run_state = resp.json()
        run_error = _extract_backend_error(run_state)
        if run_error:
            return _with_elapsed({"error": run_error}, started_at)

        try:
            state_resp = requests.get(
                f"{_url()}/threads/{thread_id}/state",
                headers={"Content-Type": "application/json"},
                timeout=20,
            )
            if state_resp.status_code == 200:
                state_json = state_resp.json()
                state_error = _extract_backend_error(state_json)
                if state_error:
                    return _with_elapsed({"error": state_error}, started_at)
                return _with_elapsed(_parse_state(state_json.get("values", state_json)), started_at)
        except Exception:
            pass

        return _with_elapsed(_parse_state(run_state), started_at)
    except requests.ConnectionError:
        return _with_elapsed({"error": "Cannot connect. Is `langgraph dev` running?"}, started_at)
    except requests.Timeout:
        return _with_elapsed({"error": "Timed out (600s). Model/backend is slow."}, started_at)
    except Exception as e:
        return _with_elapsed({"error": str(e)}, started_at)


if "langgraph_url" not in st.session_state:
    st.session_state["langgraph_url"] = LANGGRAPH_URL
if "latest_user_message" not in st.session_state:
    st.session_state["latest_user_message"] = ""
if "latest_assistant_message" not in st.session_state:
    st.session_state["latest_assistant_message"] = ""
if "latest_draft" not in st.session_state:
    st.session_state["latest_draft"] = ""
if "latest_doc_type" not in st.session_state:
    st.session_state["latest_doc_type"] = "draft"
if "latest_recommendations" not in st.session_state:
    st.session_state["latest_recommendations"] = []
if "latest_error" not in st.session_state:
    st.session_state["latest_error"] = ""
if "latest_elapsed_seconds" not in st.session_state:
    st.session_state["latest_elapsed_seconds"] = None

_inject_css()

with st.sidebar:
    st.session_state["langgraph_url"] = st.text_input("Server URL", value=_url())
    if st.button("Check Server"):
        try:
            r = requests.get(f"{_url()}/ok", timeout=5)
            st.success("Connected") if r.status_code == 200 else st.error(f"Error {r.status_code}")
        except Exception:
            st.error("Cannot connect")
    st.divider()
    if st.button("Clear Current View"):
        st.session_state["latest_user_message"] = ""
        st.session_state["latest_assistant_message"] = ""
        st.session_state["latest_draft"] = ""
        st.session_state["latest_doc_type"] = "draft"
        st.session_state["latest_recommendations"] = []
        st.session_state["latest_error"] = ""
        st.session_state["latest_elapsed_seconds"] = None
        st.rerun()
    st.caption("Stateless mode: each message uses a new thread.")

st.markdown('<div class="title">Legal Drafting Assistant</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="subtitle">Chat-style interface. No local chat history is stored.</div>',
    unsafe_allow_html=True,
)

# Render latest exchange only (no history list).
if st.session_state["latest_user_message"]:
    with st.chat_message("user"):
        st.markdown(st.session_state["latest_user_message"])

if st.session_state["latest_assistant_message"]:
    with st.chat_message("assistant"):
        st.markdown(st.session_state["latest_assistant_message"])

if st.session_state["latest_error"]:
    st.error(st.session_state["latest_error"])

if st.session_state["latest_elapsed_seconds"] is not None:
    st.caption(f"Scenario runtime: {st.session_state['latest_elapsed_seconds']:.2f}s")

if st.session_state["latest_recommendations"]:
    with st.expander("Recommendations", expanded=True):
        for rec in st.session_state["latest_recommendations"]:
            st.markdown(f"- {rec}")

draft = st.session_state["latest_draft"]
if draft:
    with st.container(border=True):
        st.subheader("Draft")
        display = re.sub(r"\{\{([A-Za-z_]+)\}\}", r"**`{{\1}}`**", draft)
        display = _format_for_markdown_display(display)
        st.markdown(display)

    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            "Download TXT",
            data=draft,
            file_name="legal_draft.txt",
            mime="text/plain",
            key="download_txt",
        )
    with col2:
        docx_bytes = _make_docx_bytes(draft)
        if docx_bytes:
            st.download_button(
                "Download DOCX",
                data=docx_bytes,
                file_name="legal_draft.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_docx",
            )
        else:
            st.caption("DOCX unavailable — install python-docx")

user_input = st.chat_input("Describe the legal document you need...")
if user_input:
    st.session_state["latest_user_message"] = user_input
    st.session_state["latest_error"] = ""

    with st.chat_message("user"):
        st.markdown(user_input)

    with st.chat_message("assistant"):
        with st.spinner("Drafting..."):
            result = send_message(user_input)

        if result.get("error"):
            st.session_state["latest_error"] = result["error"]
            st.session_state["latest_assistant_message"] = ""
            st.session_state["latest_draft"] = ""
            st.session_state["latest_elapsed_seconds"] = result.get("elapsed_seconds")
            st.error(result["error"])
        else:
            draft_text = result.get("draft", "")
            reply = _normalize_text(result.get("assistant_message", ""))
            if not reply and draft_text:
                reply = "Final draft generated. See the Draft section below."

            st.session_state["latest_assistant_message"] = reply
            st.session_state["latest_draft"] = draft_text
            st.session_state["latest_doc_type"] = result.get("doc_type", "draft")
            st.session_state["latest_recommendations"] = result.get("recommendations", [])
            st.session_state["latest_elapsed_seconds"] = result.get("elapsed_seconds")
            if reply:
                st.markdown(reply)

            # Force immediate rerun so the Draft panel (rendered above) shows this new draft now.
            if draft_text:
                st.rerun()
