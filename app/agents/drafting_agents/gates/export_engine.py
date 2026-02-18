"""
Export Engine Gate  (CLAUDE.md Step 18) -- Rule-based, NO LLM calls.

Formats the final draft into the requested export format and attaches
metadata such as word count, quality score, and generation timestamp.

Supported formats: text, docx
"""

import io
from datetime import datetime, timezone


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _extract_title(final_draft: dict) -> str:
    """Extract or derive a title from the final draft."""
    title = final_draft.get("title") or final_draft.get("heading") or ""
    if not title:
        # Fallback: use doc_type as title
        title = final_draft.get("doc_type", "Untitled Legal Document")
    return title.strip()


def _extract_sections(final_draft: dict) -> list[dict]:
    """
    Extract ordered sections from the final draft.

    Each section is expected to have: title, content (str or list[str]).
    """
    sections = final_draft.get("sections", [])
    if isinstance(sections, list):
        return sections
    return []


def _extract_prayers(final_draft: dict) -> list[str]:
    """Extract prayer / relief clauses."""
    prayers = final_draft.get("prayers", [])
    if isinstance(prayers, list):
        return [p if isinstance(p, str) else p.get("text", "") for p in prayers]
    return []


def _extract_annexures(final_draft: dict) -> list[dict]:
    """Extract annexure metadata."""
    annexures = final_draft.get("annexures", [])
    if isinstance(annexures, list):
        return annexures
    return []


def _compute_quality_score(final_draft: dict) -> float:
    """
    Derive a simple quality score from the draft metadata.

    Uses quality_score if already present; otherwise computes a
    heuristic based on completeness.
    """
    # If upstream already computed a score, use it
    existing = final_draft.get("quality_score")
    if existing is not None:
        try:
            return round(float(existing), 2)
        except (ValueError, TypeError):
            pass

    # Heuristic: award points for presence of key components
    score = 0.0
    max_score = 0.0

    checks = [
        ("title", 10),
        ("sections", 30),
        ("prayers", 20),
        ("annexures", 10),
        ("citations", 15),
        ("research", 15),
    ]

    for field, weight in checks:
        max_score += weight
        value = final_draft.get(field)
        if value:
            if isinstance(value, list) and len(value) > 0:
                score += weight
            elif isinstance(value, str) and value.strip():
                score += weight
            elif isinstance(value, dict) and len(value) > 0:
                score += weight

    if max_score == 0:
        return 0.0

    return round(score / max_score, 2)


def _format_as_text(
    title: str,
    sections: list[dict],
    prayers: list[str],
    annexures: list[dict],
) -> str:
    """
    Render the draft components into a structured plain-text document.
    """
    lines: list[str] = []

    # --- Title ---
    separator = "=" * max(len(title), 60)
    lines.append(separator)
    lines.append(title.upper())
    lines.append(separator)
    lines.append("")

    # --- Sections ---
    for idx, section in enumerate(sections, start=1):
        section_title = section.get("title", f"Section {idx}")
        content = section.get("content", "")

        lines.append(f"{idx}. {section_title}")
        lines.append("-" * (len(section_title) + len(str(idx)) + 2))

        if isinstance(content, list):
            for item in content:
                if isinstance(item, str):
                    lines.append(f"   {item}")
                elif isinstance(item, dict):
                    lines.append(f"   {item.get('text', str(item))}")
        elif isinstance(content, str):
            # Indent each line of content
            for line in content.split("\n"):
                lines.append(f"   {line}")
        else:
            lines.append(f"   {content}")

        lines.append("")

    # --- Prayers ---
    if prayers:
        lines.append("PRAYER / RELIEF SOUGHT")
        lines.append("-" * 22)
        for pidx, prayer in enumerate(prayers, start=1):
            lines.append(f"   {pidx}. {prayer}")
        lines.append("")

    # --- Annexures ---
    if annexures:
        lines.append("ANNEXURES")
        lines.append("-" * 10)
        for aidx, annexure in enumerate(annexures, start=1):
            ann_title = annexure.get("title", f"Annexure {aidx}")
            ann_desc = annexure.get("description", "")
            lines.append(f"   Annexure-{aidx}: {ann_title}")
            if ann_desc:
                lines.append(f"      {ann_desc}")
        lines.append("")

    return "\n".join(lines)


def _format_as_docx(
    title: str,
    sections: list[dict],
    prayers: list[str],
    annexures: list[dict],
) -> bytes:
    """
    Render the draft components into a DOCX document (bytes).

    Returns the DOCX file content as bytes for storage or download.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)

    # --- Title ---
    title_para = doc.add_heading(title.upper(), level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # spacer

    # --- Sections ---
    for idx, section in enumerate(sections, start=1):
        section_title = section.get("title", f"Section {idx}")
        content = section.get("content", "")

        heading = doc.add_heading(f"{idx}. {section_title}", level=2)

        if isinstance(content, list):
            for item in content:
                text = item if isinstance(item, str) else item.get("text", str(item))
                p = doc.add_paragraph(text)
                p.paragraph_format.left_indent = Inches(0.5)
        elif isinstance(content, str):
            for line in content.split("\n"):
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    p.paragraph_format.left_indent = Inches(0.5)
        else:
            doc.add_paragraph(str(content))

    # --- Prayers ---
    if prayers:
        doc.add_heading("PRAYER / RELIEF SOUGHT", level=1)
        for pidx, prayer in enumerate(prayers, start=1):
            p = doc.add_paragraph(f"{pidx}. {prayer}")
            p.paragraph_format.left_indent = Inches(0.5)

    # --- Annexures ---
    if annexures:
        doc.add_heading("ANNEXURES", level=1)
        for aidx, annexure in enumerate(annexures, start=1):
            ann_title = annexure.get("title", f"Annexure {aidx}")
            ann_desc = annexure.get("description", "")
            p = doc.add_paragraph(f"Annexure-{aidx}: {ann_title}")
            p.runs[0].bold = True
            if ann_desc:
                doc.add_paragraph(f"   {ann_desc}")

    # Serialize to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def prepare_export(
    final_draft: dict,
    format: str = "text",
) -> dict:
    """
    Export Engine gate (Step 18).

    Formats the final draft for export.  Pure rule-based -- no LLM calls.

    Supported formats: ``"text"``, ``"docx"``.

    Args:
        final_draft: The completed draft dict with keys like title,
                     sections, prayers, annexures, citations, etc.
        format:      Export format (default ``"text"``).

    Returns:
        dict with keys:
            gate           - "export_engine"
            passed         - True
            export_format  - str
            export_content - str (text) or bytes (docx)
            metadata       - {generated_at, quality_score, word_count,
                              section_count, has_prayers, has_annexures}
    """
    supported_formats = {"text", "docx"}

    if format not in supported_formats:
        return {
            "gate": "export_engine",
            "passed": False,
            "export_format": format,
            "export_content": "",
            "metadata": {
                "error": f"Unsupported format '{format}'. "
                         f"Supported: {', '.join(sorted(supported_formats))}.",
            },
        }

    # Extract components
    title = _extract_title(final_draft)
    sections = _extract_sections(final_draft)
    prayers = _extract_prayers(final_draft)
    annexures = _extract_annexures(final_draft)

    # Format
    export_content = None
    if format == "text":
        export_content = _format_as_text(title, sections, prayers, annexures)
    elif format == "docx":
        export_content = _format_as_docx(title, sections, prayers, annexures)

    # Metadata (word count uses text rendering for all formats)
    text_content = _format_as_text(title, sections, prayers, annexures) if format != "text" else export_content
    word_count = len(text_content.split()) if text_content else 0
    quality_score = _compute_quality_score(final_draft)

    metadata = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "quality_score": quality_score,
        "word_count": word_count,
        "section_count": len(sections),
        "has_prayers": len(prayers) > 0,
        "has_annexures": len(annexures) > 0,
    }

    return {
        "gate": "export_engine",
        "passed": True,
        "export_format": format,
        "export_content": export_content,
        "metadata": metadata,
    }
